import streamlit as st
import tensorflow as tf
from PIL import Image
import numpy as np
import base64
import os
import matplotlib.pyplot as plt
import cv2
from datetime import datetime
import json
from docx import Document
import io

# ========================================
# 📱 PAGE CONFIGURATION - إعدادات الصفحة
# ========================================
st.set_page_config(
    page_title="Medical AI - Chest X-Ray Diagnostic",
    page_icon="🩺",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========================================
# 🎨 CSS STYLING - تحسين الواجهة
# ========================================
css = """
<style>
/* الخلفية الرئيسية */
.stApp {
    background: linear-gradient(135deg, #0f1419 0%, #1a1f2e 100%);
}

/* العنوان */
h1 {
    color: #00d4ff !important;
    text-shadow: 2px 2px 10px rgba(0,0,0,0.7) !important;
    font-weight: 900 !important;
}

h2, h3 {
    color: #00d4ff !important;
}

/* صندوق النتيجة */
.prediction-box {
    padding: 30px;
    border-radius: 15px;
    text-align: center;
    color: white;
    font-size: 28px;
    font-weight: bold;
    box-shadow: 0 8px 16px rgba(0,0,0,0.4);
    border: 2px solid #00d4ff;
}

/* بطاقات المعلومات */
.info-card {
    background: rgba(0, 212, 255, 0.1);
    border-left: 5px solid #00d4ff;
    padding: 15px;
    border-radius: 8px;
    color: #fff;
    margin: 10px 0;
}

/* النص */
p, span {
    color: #ccc !important;
}

/* الأزرار */
.stButton > button {
    background-color: #00d4ff !important;
    color: #000 !important;
    font-weight: bold !important;
    border-radius: 8px !important;
    padding: 10px 30px !important;
    border: none !important;
}

.stButton > button:hover {
    background-color: #00b8d4 !important;
    transform: scale(1.05);
}

/* التحذيرات والمعلومات */
.stAlert {
    background-color: rgba(255, 107, 107, 0.2) !important;
    color: #fff !important;
}
</style>
"""

st.markdown(css, unsafe_allow_html=True)

# ========================================
# 🏥 الرأس الرئيسي (Header)
# ========================================
col_uni, col_title, col_dept = st.columns([1, 3, 1])

with col_uni:
    st.image("https://uomus.edu.iq/img/logo.png", width=100) 

with col_title:
    st.markdown("""
    <div style='text-align: center;'>
        <h1 style='margin-bottom: 0; font-size: 2rem;'>🩺 Smart Chest X-Ray Diagnostic System</h1>
        <p style='color: #00d4ff; font-size: 1.1rem; margin-top: 5px;'>Intelligent Medical Image Analysis</p>
        <p style='color: #888; font-size: 0.9rem;'>Al-Mustaqbal University</p>
    </div>
    """, unsafe_allow_html=True)

with col_dept:
    try:
        st.image("/content/logo2.png", width=100)
    except:
        st.markdown("<p style='text-align: center; font-size: 0.8rem; color: #00d4ff;'>Department Logo</p>", unsafe_allow_html=True)

st.markdown("---")

# ========================================
# 🔧 تحميل النموذج المدرب
# ========================================
@st.cache_resource
def load_trained_model():
    """تحميل النموذج المدرب"""
    try:
        model_paths = ['medical_multi_model.h5', '/content/medical_multi_model.h5', 'best_model.h5']
        for path in model_paths:
            if os.path.exists(path):
                try:
                    # محاولة التحميل العادي
                    return tf.keras.models.load_model(path), True
                except Exception as e:
                    # إذا فشل، حمّل مع الإعدادات القديمة
                    st.warning(f"⚠️ محاولة تحميل النموذج بطريقة متوافقة...")
                    try:
                        model = tf.keras.models.load_model(
                            path, 
                            custom_objects=None,
                            compile=False  # تحميل بدون تجميع
                        )
                        model.compile()
                        return model, True
                    except Exception as e2:
                        st.error(f"❌ فشل التحميل: {str(e2)}")
                        return None, False
        return None, False
    except Exception as e:
        st.error(f"❌ خطأ في تحميل النموذج: {str(e)}")
        return None, False

# ========================================
# 🎯 الشريط الجانبي
# ========================================
st.sidebar.markdown("### ⚙️ الإعدادات")
confidence_threshold = st.sidebar.slider("📌 حد الثقة الأدنى (%)", 0, 100, 50, 5)
display_mode = st.sidebar.radio("🎯 وضع العرض", ["بسيط", "متقدم", "تقرير"])

st.sidebar.markdown("---")
st.sidebar.markdown("### 📚 أنواع الأمراض")
for disease in class_names:
    st.sidebar.info(f"**{disease}**")

# ========================================
# 📤 رفع الصورة
# ========================================
st.markdown("### 📤 رفع صورة الأشعة السينية")
uploaded_file = st.file_uploader("اختر صورة (JPG, PNG, JPEG)", type=["jpg", "png", "jpeg"])

if uploaded_file:
    img = Image.open(uploaded_file).convert('RGB')
    
    # التنبؤ
    img_resized = img.resize((224, 224))
    img_norm = np.array(img_resized) / 255.0
    img_expanded = np.expand_dims(img_norm, axis=0)
    predictions = model.predict(img_expanded, verbose=0)
    
    predicted_idx = np.argmax(predictions[0])
    label = class_names[predicted_idx]
    color = colors.get(label, "#3498db")
    confidence_val = np.max(predictions[0]) * 100

    # ========================================
    # 📊 عرض النتائج
    # ========================================
    if display_mode == "بسيط":
        col1, col2 = st.columns(2)
        with col1:
            st.image(img, caption='الصورة المرفوعة', use_container_width=True)
        with col2:
            st.markdown("### 🔬 النتيجة")
            st.markdown(f"""
            <div class='prediction-box' style='background: {color}33; border-color: {color};'>
                <div>{label}</div>
                <small>الثقة: {confidence_val:.1f}%</small>
            </div>
            """, unsafe_allow_html=True)
            
            if confidence_val < confidence_threshold:
                st.warning(f"⚠️ الثقة أقل من {confidence_threshold}%")
    
    elif display_mode == "متقدم":
        st.markdown("### 📊 التحليل المتقدم")
        col1, col2 = st.columns([1.2, 1])
        
        with col1:
            st.image(img, caption='الصورة المرفوعة', use_container_width=True)
        
        with col2:
            st.markdown(f"""
            <div class='prediction-box' style='background: {color}33; border-color: {color};'>
                {label}<br>
                <small>{confidence_val:.2f}%</small>
            </div>
            """, unsafe_allow_html=True)
            
            # رسم بياني
            fig, ax = plt.subplots(figsize=(10, 5))
            probs = predictions[0] * 100
            colors_list = [colors.get(d, '#3498db') for d in class_names]
            
            bars = ax.barh(class_names, probs, color=colors_list)
            ax.set_xlabel('الثقة (%)', color='white')
            ax.set_facecolor('#1a1a1a')
            fig.patch.set_facecolor('#1a1a1a')
            ax.tick_params(colors='white')
            
            for bar, prob in zip(bars, probs):
                ax.text(prob + 1, bar.get_y() + bar.get_height()/2, f'{prob:.1f}%', 
                       va='center', color='white', fontweight='bold')
            
            st.pyplot(fig)

    elif display_mode == "تقرير":
        st.markdown("### 📋 التقرير الطبي الكامل")
        
        # عرض صورة مصغرة
        st.image(img, width=250, caption="صورة الأشعة المرفقة")
        st.success(f"التشخيص النهائي: {label}")

        def create_word_report():
            doc = Document()
            doc.add_heading('Medical AI Diagnostic Report', 0)
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Engineer: ABDULLAH BSHAR SALIH")
            
            doc.add_heading('Analysis Results', level=1)
            p = doc.add_paragraph()
            p.add_run('Diagnosis: ').bold = True
            p.add_run(label)
            
            p = doc.add_paragraph()
            p.add_run('Confidence Level: ').bold = True
            p.add_run(f"{confidence_val:.2f}%")
            
            doc.add_heading('Technical Details', level=2)
            for i, name in enumerate(class_names):
                doc.add_paragraph(f"- {name}: {predictions[0][i]*100:.2f}%")
            
            doc.add_heading('Academic Reference', level=2)
            doc.add_paragraph("AI Engineering Department - Al-Mustaqbal University.")
            
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            return bio.getvalue()

        word_data = create_word_report()
        st.download_button(
            label="📄 تحميل التقرير بصيغة Word",
            data=word_data,
            file_name=f"Medical_Report_{label}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ========================================
# 👨‍💼 الفريق
# ========================================
st.markdown("---")
st.markdown("""
<div class='info-card'>
    <h3>👨‍💼 فريق التطوير</h3>
    <p><strong>المهندس:</strong> ABDULLAH BSHAR SALIH</p>
    <p><strong>التخصص:</strong> هندسة الذكاء الاصطناعي</p>
    <p><strong>الجامعة:</strong> جامعة المستقبل</p>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div style='background: #ff6b6b20; border-left: 4px solid #ff6b6b; padding: 15px; border-radius: 8px; color: white;'>
    <h4>⚠️ تنبيه مهم</h4>
    <p>هذا النظام <strong>للمساعدة الأولية فقط</strong> وليس بديل عن الطبيب المتخصص</p>
</div>
""", unsafe_allow_html=True)

print("✅ تطبيق Streamlit جاهز!")
